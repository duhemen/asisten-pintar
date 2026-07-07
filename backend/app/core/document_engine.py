import os
import re
from typing import Optional, Dict, Any
import logging

# Import dengan try-except untuk menghindari error jika library tidak terinstal
try:
    import docx
except ImportError:
    docx = None
    print("Warning: python-docx tidak terinstal. Docx tidak akan didukung.")

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    print("Warning: PyPDF2 tidak terinstal. PDF tidak akan didukung.")

try:
    import pandas as pd
except ImportError:
    pd = None
    print("Warning: pandas tidak terinstal. Excel tidak akan didukung.")

try:
    from openpyxl import load_workbook
except ImportError:
    load_workbook = None
    print("Warning: openpyxl tidak terinstal. Excel tidak akan didukung.")

from app.core.ocr_engine import ocr_processor

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class LocalDocumentEngine:
    def __init__(self):
        self.supported_extensions = {
            '.txt': self._extract_txt,
            '.docx': self._extract_docx,
            '.pdf': self._extract_pdf,
            '.xlsx': self._extract_xlsx,
            '.xls': self._extract_xlsx,
            '.jpg': self._extract_image,
            '.jpeg': self._extract_image,
            '.png': self._extract_image,
        }
        
        # Cek library yang tersedia
        self._check_libraries()
    
    def _check_libraries(self):
        """Cek ketersediaan library"""
        if docx is None:
            logger.warning("python-docx tidak tersedia. File .docx tidak akan didukung.")
        if PyPDF2 is None:
            logger.warning("PyPDF2 tidak tersedia. File .pdf tidak akan didukung.")
        if pd is None and load_workbook is None:
            logger.warning("pandas/openpyxl tidak tersedia. File .xlsx tidak akan didukung.")
    
    def extract_text(self, file_path: str, filename: str, use_vision: bool = True) -> Dict[str, Any]:
        """
        Ekstrak teks dari berbagai format file dengan deteksi PDF scan otomatis
        
        Returns:
            dict: {
                'text': str,
                'method': str,
                'pages': int,
                'confidence': float
            }
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File tidak ditemukan: {file_path}")
        
        ext = os.path.splitext(filename)[1].lower()
        
        if ext not in self.supported_extensions:
            raise ValueError(f"Format file tidak didukung: {ext}")
        
        # Cek library untuk format tertentu
        if ext == '.docx' and docx is None:
            raise ValueError("python-docx tidak terinstal. Tidak dapat membaca file .docx")
        if ext == '.pdf' and PyPDF2 is None:
            raise ValueError("PyPDF2 tidak terinstal. Tidak dapat membaca file .pdf")
        if ext in ['.xlsx', '.xls'] and (pd is None and load_workbook is None):
            raise ValueError("pandas/openpyxl tidak terinstal. Tidak dapat membaca file Excel")
        
        try:
            result = self.supported_extensions[ext](file_path, use_vision)
            return result
        except Exception as e:
            logger.error(f"Error ekstraksi {filename}: {e}")
            raise
    
    def _extract_txt(self, file_path: str, use_vision: bool = True) -> Dict[str, Any]:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return {
            'text': text,
            'method': 'txt',
            'pages': 1,
            'confidence': 1.0
        }
    
    def _extract_docx(self, file_path: str, use_vision: bool = True) -> Dict[str, Any]:
        if docx is None:
            raise ImportError("python-docx tidak terinstal")
        
        doc = docx.Document(file_path)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        return {
            'text': '\n'.join(full_text),
            'method': 'docx',
            'pages': len(doc.paragraphs) // 10 + 1,
            'confidence': 1.0
        }
    
    def _extract_pdf(self, file_path: str, use_vision: bool = True) -> Dict[str, Any]:
        """Ekstrak teks dari PDF dengan deteksi scan otomatis"""
        if PyPDF2 is None:
            raise ImportError("PyPDF2 tidak terinstal")
        
        # Coba ekstrak teks normal dulu
        normal_text = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                total_pages = len(pdf_reader.pages)
                text_pages = 0
                
                for page_num in range(total_pages):
                    try:
                        page = pdf_reader.pages[page_num]
                        page_text = page.extract_text()
                        if page_text and page_text.strip():
                            normal_text += f"Halaman {page_num + 1}:\n{page_text}\n\n"
                            text_pages += 1
                    except Exception as e:
                        logger.warning(f"Gagal ekstrak halaman {page_num + 1}: {e}")
                
                # Jika lebih dari 50% halaman ada teks, gunakan hasil normal
                if total_pages > 0 and text_pages > total_pages * 0.5:
                    return {
                        'text': normal_text,
                        'method': 'pdf_text',
                        'pages': total_pages,
                        'confidence': 1.0
                    }
        except Exception as e:
            logger.warning(f"Gagal ekstrak normal: {e}")
        
        # Jika tidak, gunakan OCR/Vision
        logger.info("PDF terdeteksi sebagai scan, menggunakan OCR/Vision")
        try:
            result = ocr_processor.extract_text_from_pdf(file_path, use_vision=use_vision)
            return result
        except Exception as e:
            logger.error(f"OCR/Vision gagal: {e}")
            # Return error dengan pesan yang jelas
            return {
                'text': '',
                'method': 'error',
                'pages': 0,
                'confidence': 0
            }
    
    def _extract_xlsx(self, file_path: str, use_vision: bool = True) -> Dict[str, Any]:
        """Ekstrak teks dari Excel dengan handling multiple sheets"""
        try:
            # Coba dengan pandas dulu
            if pd is not None:
                xlsx = pd.ExcelFile(file_path)
                all_text = []
                total_sheets = len(xlsx.sheet_names)
                
                for sheet_name in xlsx.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    all_text.append(f"--- Sheet: {sheet_name} ---")
                    # Konversi ke teks
                    text = df.to_string(index=False)
                    all_text.append(text)
                    all_text.append("")
                
                return {
                    'text': '\n'.join(all_text),
                    'method': 'excel_pandas',
                    'pages': total_sheets,
                    'confidence': 1.0
                }
            elif load_workbook is not None:
                # Fallback ke openpyxl
                workbook = load_workbook(file_path, data_only=True)
                all_text = []
                total_sheets = len(workbook.sheetnames)
                
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    all_text.append(f"--- Sheet: {sheet_name} ---")
                    for row in sheet.iter_rows(values_only=True):
                        row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
                        if row_text.strip():
                            all_text.append(row_text)
                    all_text.append("")
                
                return {
                    'text': '\n'.join(all_text),
                    'method': 'excel_openpyxl',
                    'pages': total_sheets,
                    'confidence': 1.0
                }
            else:
                raise ImportError("pandas atau openpyxl tidak terinstal")
                
        except Exception as e:
            raise ValueError(f"Gagal membaca Excel: {str(e)}")
    
    def _extract_image(self, file_path: str, use_vision: bool = True) -> Dict[str, Any]:
        """Ekstrak teks dari gambar"""
        result = ocr_processor.extract_text_from_image(file_path, use_vision=use_vision)
        return result

# Singleton
local_doc_engine = LocalDocumentEngine()