# backend/app/modules/tnd_validator/service.py
import re
import os
import logging
from typing import Dict, List, Optional, Any
from datetime import datetime
from typing import Dict, List, Optional, Any

# Import untuk membaca dokumen
try:
    from docx import Document
except ImportError:
    Document = None
    print("Warning: python-docx tidak terinstal")

try:
    from pptx import Presentation
except ImportError:
    Presentation = None
    print("Warning: python-pptx tidak terinstal")

try:
    import PyPDF2
except ImportError:
    PyPDF2 = None
    print("Warning: PyPDF2 tidak terinstal")

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class TNDValidator:
    """
    Validator Tata Naskah Dinas (TND)
    Memeriksa kesesuaian format surat dinas sesuai peraturan
    """
    
    def __init__(self):
        # Standar TND berdasarkan aturan umum
        self.standar = {
            "font": {
                "nama": ["Arial", "Times New Roman", "Calibri"],
                "ukuran": 12,
                "spasi": 1.5
            },
            "margin": {
                "top": 4,    # cm
                "bottom": 3,  # cm
                "left": 4,    # cm
                "right": 3    # cm
            },
            "nomor_surat": r"\d{3}/\d+/([A-Z]+\.?)+/[A-Z]+/\d{4}",
            "kop_surat": ["PEMERINTAH", "KABUPATEN", "KOTA", "PROVINSI"],
            "tanda_tangan": ["ttd", "mengetahui", "an.", "a.n.", "atas nama"]
        }
        
        self.hasil = {
            "is_valid": True,
            "errors": [],
            "warnings": [],
            "suggestions": [],
            "details": {}
        }
    
    def validate_file(self, file_path: str, filename: str) -> Dict[str, Any]:
        """
        Validasi dokumen berdasarkan format file
        """
        ext = os.path.splitext(filename)[1].lower()
        
        # Reset hasil
        self.hasil = {
            "is_valid": True,
            "errors": [],
            "warnings": [],
            "suggestions": [],
            "details": {},
            "filename": filename,
            "timestamp": datetime.now().isoformat()
        }
        
        # Pilih metode berdasarkan ekstensi
        if ext == '.docx':
            return self._validate_docx(file_path)
        elif ext == '.pdf':
            return self._validate_pdf(file_path)
        elif ext == '.txt':
            return self._validate_txt(file_path)
        else:
            self.hasil["errors"].append(f"Format {ext} tidak didukung untuk validasi TND")
            self.hasil["is_valid"] = False
            return self.hasil
    
    def _validate_docx(self, file_path: str) -> Dict[str, Any]:
        """Validasi file DOCX"""
        if Document is None:
            self.hasil["errors"].append("python-docx tidak terinstal")
            self.hasil["is_valid"] = False
            return self.hasil
        
        try:
            doc = Document(file_path)
            
            # Ekstrak teks
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            text_content = "\n".join(full_text)
            
            # Cek properti dokumen
            self._check_document_properties(text_content, doc)
            
            # Cek format nomor surat
            self._check_nomor_surat(text_content)
            
            # Cek kop surat
            self._check_kop_surat(text_content)
            
            # Cek tanda tangan
            self._check_tanda_tangan(text_content)
            
            # Cek struktur surat
            self._check_struktur_surat(full_text)
            
            # Simpan teks lengkap
            self.hasil["details"]["full_text"] = text_content
            self.hasil["details"]["word_count"] = len(text_content.split())
            
            return self.hasil
            
        except Exception as e:
            logger.error(f"Error validasi DOCX: {e}")
            self.hasil["errors"].append(f"Gagal membaca dokumen: {str(e)}")
            self.hasil["is_valid"] = False
            return self.hasil
    
    def _validate_pdf(self, file_path: str) -> Dict[str, Any]:
        """Validasi file PDF"""
        if PyPDF2 is None:
            self.hasil["errors"].append("PyPDF2 tidak terinstal")
            self.hasil["is_valid"] = False
            return self.hasil
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                text_content = ""
                
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            # Cek format nomor surat
            self._check_nomor_surat(text_content)
            
            # Cek kop surat
            self._check_kop_surat(text_content)
            
            # Cek tanda tangan
            self._check_tanda_tangan(text_content)
            
            # Simpan teks
            self.hasil["details"]["full_text"] = text_content
            self.hasil["details"]["pages"] = len(pdf_reader.pages)
            
            return self.hasil
            
        except Exception as e:
            logger.error(f"Error validasi PDF: {e}")
            self.hasil["errors"].append(f"Gagal membaca PDF: {str(e)}")
            self.hasil["is_valid"] = False
            return self.hasil
    
    def _validate_txt(self, file_path: str) -> Dict[str, Any]:
        """Validasi file TXT"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text_content = f.read()
            
            # Cek format nomor surat
            self._check_nomor_surat(text_content)
            
            # Cek kop surat
            self._check_kop_surat(text_content)
            
            # Cek tanda tangan
            self._check_tanda_tangan(text_content)
            
            # Simpan teks
            self.hasil["details"]["full_text"] = text_content
            self.hasil["details"]["word_count"] = len(text_content.split())
            
            return self.hasil
            
        except Exception as e:
            logger.error(f"Error validasi TXT: {e}")
            self.hasil["errors"].append(f"Gagal membaca file: {str(e)}")
            self.hasil["is_valid"] = False
            return self.hasil
    
    def _check_document_properties(self, text: str, doc) -> None:
        """Cek properti dokumen (font, ukuran, dll)"""
        try:
            # Cek font pada paragraf pertama
            if doc.paragraphs and doc.paragraphs[0].runs:
                run = doc.paragraphs[0].runs[0]
                
                # Cek font name
                if run.font.name:
                    font_name = run.font.name
                    if font_name not in self.standar["font"]["nama"]:
                        self.hasil["warnings"].append(
                            f"Font '{font_name}' tidak sesuai standar TND. Gunakan: {', '.join(self.standar['font']['nama'])}"
                        )
                    else:
                        self.hasil["details"]["font_name"] = font_name
                
                # Cek font size
                if run.font.size:
                    font_size = run.font.size.pt
                    if font_size != self.standar["font"]["ukuran"]:
                        self.hasil["warnings"].append(
                            f"Ukuran font {font_size}pt tidak sesuai standar TND (harus {self.standar['font']['ukuran']}pt)"
                        )
                    else:
                        self.hasil["details"]["font_size"] = font_size
            
        except Exception as e:
            logger.warning(f"Gagal cek properti dokumen: {e}")
    
    def _check_nomor_surat(self, text: str) -> None:
        """Cek format nomor surat"""
        pattern = self.standar["nomor_surat"]
        matches = re.findall(pattern, text)
        
        if matches:
            self.hasil["details"]["nomor_surat"] = matches[0] if matches else None
            self.hasil["details"]["nomor_surat_count"] = len(matches)
        else:
            self.hasil["errors"].append(
                "Format Nomor Surat tidak terdeteksi. Contoh: 005/123/Diskominfo/A/2024"
            )
            self.hasil["is_valid"] = False
            self.hasil["suggestions"].append(
                "Pastikan nomor surat sesuai format: 000/000/Instansi/Kode/Bulan/Tahun"
            )
    
    def _check_kop_surat(self, text: str) -> None:
        """Cek keberadaan kop surat"""
        text_upper = text.upper()
        found = False
        found_keywords = []
        
        for keyword in self.standar["kop_surat"]:
            if keyword in text_upper:
                found = True
                found_keywords.append(keyword)
        
        if found:
            self.hasil["details"]["kop_surat"] = found_keywords
        else:
            self.hasil["warnings"].append(
                "Kop surat tidak terdeteksi. Pastikan surat menggunakan kop instansi resmi"
            )
    
    def _check_tanda_tangan(self, text: str) -> None:
        """Cek tanda tangan/penandatangan"""
        text_lower = text.lower()
        found = False
        found_keywords = []
        
        for keyword in self.standar["tanda_tangan"]:
            if keyword in text_lower:
                found = True
                found_keywords.append(keyword)
        
        if found:
            self.hasil["details"]["tanda_tangan"] = found_keywords
        else:
            self.hasil["warnings"].append(
                "Tanda tangan/penandatangan tidak terdeteksi"
            )
            self.hasil["suggestions"].append(
                "Tambahkan bagian tanda tangan dengan format: (Nama Lengkap)\nNIP."
            )
    
    def _check_struktur_surat(self, paragraphs: List[str]) -> None:
        """Cek struktur dasar surat"""
        if len(paragraphs) < 3:
            self.hasil["warnings"].append(
                "Struktur surat terlihat pendek. Pastikan memiliki: Kepala Surat → Isi → Penutup"
            )
            return
        
        # Cek kata kunci struktur
        struktur_check = {
            "pembuka": ["dengan", "berdasarkan", "sehubungan", "menunjuk"],
            "isi": ["bahwa", "untuk", "agar", "mohon", "dapat"],
            "penutup": ["demikian", "atas", "terima kasih", "diberitahukan"]
        }
        
        text = " ".join(paragraphs).lower()
        
        for bagian, keywords in struktur_check.items():
            found = any(kw in text for kw in keywords)
            if not found:
                self.hasil["warnings"].append(
                    f"Bagian {bagian} surat tidak terdeteksi dengan jelas"
                )
    
    def get_summary(self) -> Dict[str, Any]:
        """Dapatkan ringkasan hasil validasi"""
        return {
            "is_valid": self.hasil["is_valid"],
            "status": "✅ Valid" if self.hasil["is_valid"] else "❌ Perlu Perbaikan",
            "total_errors": len(self.hasil["errors"]),
            "total_warnings": len(self.hasil["warnings"]),
            "total_suggestions": len(self.hasil["suggestions"]),
            "errors": self.hasil["errors"],
            "warnings": self.hasil["warnings"],
            "suggestions": self.hasil["suggestions"],
            "details": self.hasil.get("details", {})
        }


# Singleton validator
tnd_validator = TNDValidator()